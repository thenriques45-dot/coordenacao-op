# -*- coding: utf-8 -*-
"""Gera um PLANO DE ENSINO de exemplo (parser + gerador) com dados reais da planilha."""
import re
from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REFERENCIAS = ["Currículo Priorizado", "Escopo Sequência", "Currículo Paulista", "BNCC"]
LARGURA_TOTAL = 9360  # twips (~6,5") — mesma para todas as tabelas (alinhamento)

# ── Parser (mesma lógica a ser portada para o Rust) ──────────────────────────
COD_BNCC = re.compile(r"E[FM]\d{2}[A-Z]{2,3}\d{2,3}[A-Z*]*")

def parse_aula(s):
    """Retorna (unidade, objetos, habilidade), detectando formato EM ou EF."""
    m = re.match(r"\s*Aula\s+\d+\s*[—–-]\s*(.*)", s, re.S)
    resto = m.group(1) if m else s
    segs = [p.strip() for p in resto.split("|")]
    unidade, objetos, habilidade = segs[0].strip(), "", ""
    rotulado = any(seg.lower().startswith(("conteúdos", "conteudos", "habilidades")) for seg in segs[1:])
    if rotulado:  # formato EF (campos rotulados)
        for seg in segs[1:]:
            low = seg.lower()
            if low.startswith(("conteúdos", "conteudos")):
                objetos = seg.split(":", 1)[1].strip() if ":" in seg else seg
            elif low.startswith("habilidades"):
                habilidade = seg.split(":", 1)[1].strip() if ":" in seg else seg
        if not habilidade:
            habilidade = ", ".join(dict.fromkeys(COD_BNCC.findall(s)))
    else:  # formato EM:  {objetos} (AE.. | EM13..)
        rest = " | ".join(segs[1:]) if len(segs) > 1 else ""
        pm = re.search(r"\(([^()]*)\)\s*$", rest)
        if pm:
            habilidade, objetos = pm.group(1).strip(), rest[:pm.start()].strip()
        else:
            objetos = rest
    return unidade, objetos, habilidade

def dedup(seq):
    return list(dict.fromkeys(x for x in seq if x))

def turmas_legivel(resposta):
    """'Turma A, Turma B' -> 'A, B'."""
    letras = [t.strip().replace("Turma", "").strip() for t in resposta.split(",")]
    return ", ".join(l for l in letras if l)

# ── Helpers docx ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, cor):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), cor)
    cell._tc.get_or_add_tcPr().append(shd)

def fixar_larguras(tabela, larguras):
    """Layout fixo + largura por célula (alinhamento consistente)."""
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False
    tblPr = tabela._tbl.tblPr
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblPr.append(layout)
    w = OxmlElement("w:tblW"); w.set(qn("w:w"), str(sum(larguras))); w.set(qn("w:type"), "dxa"); tblPr.append(w)
    for row in tabela.rows:
        for cell, larg in zip(row.cells, larguras):
            cell.width = Twips(larg)

def run(p, txt, bold=False, size=10, italic=False):
    r = p.add_run(txt); r.bold = bold; r.italic = italic; r.font.size = Pt(size); return r

def bullets(cell, itens):
    """itens: lista de strings -> cada uma um tópico."""
    cell.text = ""
    if not itens:
        itens = [""]
    for i, l in enumerate(itens):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run(p, ("• " + l) if l else "", size=10)

def split_itens(texto):
    """Quebra 'a; b.; c' em ['a','b','c'] (para conteúdos/estratégias)."""
    return [l.strip(" ;.").strip() for l in re.split(r";|\n", texto) if l.strip(" ;.")]

# ── Geração ──────────────────────────────────────────────────────────────────
def gerar(caminho, d):
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Twips(720)

    # Cabeçalho: imagem configurada no programa (placeholder no exemplo)
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; fixar_larguras(t, [LARGURA_TOTAL])
    c = t.rows[0].cells[0]; set_cell_bg(c, "F2F2F2")
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "[ CABEÇALHO — imagem configurada em Configurações → Instituição ]", italic=True, size=9)
    doc.add_paragraph()

    # Título + bimestre
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; fixar_larguras(t, [LARGURA_TOTAL])
    p = t.rows[0].cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, f"PLANO DE ENSINO\n{d['bimestre']}º Bimestre", bold=True, size=12)

    # Professor
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; fixar_larguras(t, [LARGURA_TOTAL])
    p = t.rows[0].cells[0].paragraphs[0]
    run(p, "Professor: ", bold=True); run(p, d["professor"])

    # Identificação: Disciplina | Série/ano | Turmas | Ano letivo
    cols = [("Disciplina: ", d["disciplina"]), ("Série/ano: ", d["ano"]),
            ("Turmas: ", d["turmas"]), ("Ano letivo: ", d["ano_letivo"])]
    t = doc.add_table(rows=1, cols=len(cols)); t.style = "Table Grid"
    fixar_larguras(t, [3000, 2400, 2400, 1560])
    for cell, (lab, val) in zip(t.rows[0].cells, cols):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, lab, bold=True); run(p, val)

    # Cabeçalho do bimestre (cinza)
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; fixar_larguras(t, [LARGURA_TOTAL])
    c = t.rows[0].cells[0]; set_cell_bg(c, "D9D9D9")
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, f"{d['bimestre']}º BIMESTRE", bold=True, size=11)

    # Linhas rótulo | conteúdo
    linhas = [
        ("UNIDADE TEMÁTICA", d["unidade"]),
        ("OBJETOS DE CONHECIMENTO", d["objetos"]),
        ("HABILIDADE", d["habilidade"]),
        ("ESTRATÉGIAS / METODOLOGIA", split_itens(d["estrategias"])),
        ("RECURSOS PEDAGÓGICOS", split_itens(d["recursos"])),
        ("AVALIAÇÃO", split_itens(d["avaliacao"])),
        ("ADAPTAÇÃO CURRICULAR", split_itens(d["adaptacao"])),
        ("COMO VERIFICAR SE O OBJETIVO FOI CUMPRIDO", split_itens(d["verificacao"])),
        ("REFERÊNCIAS", REFERENCIAS),
    ]
    t = doc.add_table(rows=len(linhas), cols=2); t.style = "Table Grid"
    fixar_larguras(t, [2800, LARGURA_TOTAL - 2800])
    for row, (lab, itens) in zip(t.rows, linhas):
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        run(row.cells[0].paragraphs[0], lab, bold=True)
        bullets(row.cells[1], itens)

    doc.save(caminho); print("Documento salvo em:", caminho)

# ── Dados REAIS da planilha (1ª resposta) ────────────────────────────────────
AULAS = [
    "Aula 1 - Descobrindo as lutas do mundo | Conteúdos: Conceito de lutas do mundo.; Origem do Judô, Jiu-Jitsu e do Caratê.; Características específicas do Judô, Jiu- | Habilidades: EF08EF16, EF08EF17 | AE3 - Vivenciar as lutas do mundo, explorando técnicas e golpes de diferentes modalidades, respeitando os colegas.",
    "Aula 7 - Saudações e bases do caratê | Conteúdos: Saudações e bases no caratê.; Vivenciar as saudações reconhecendo os valores relacionados ao respeito.; Experimentar as posturas e bases do caratê. | Habilidades: EF08EF16, EF08EF17 | AE3 - Vivenciar as lutas do mundo, explorando técnicas e golpes de diferentes modalidades, respeitando os colegas.",
    "Aula 11 - Esporte paralímpico: judô para cegos | Conteúdos: Origem e evolução do Judô para cegos.; Mudanças e adequações às regras.; Conhecer a origem e a evolução do judô para cego. | Habilidades: EF08EF21* | AE4 - Experimentar o judô para cegos, analisando suas particularidades técnicas e táticas, discutindo estereótipos e preconceitos.",
]
parsed = [parse_aula(a) for a in AULAS]

dados = {
    "professor": "Elisane Brum Coelho",
    "disciplina": "Educação Física",
    "ano": "8º Ano",
    "turmas": turmas_legivel("Turma A, Turma B, Turma C, Turma D, Turma E, Turma F, Turma G"),
    "ano_letivo": str(date.today().year),  # sempre o ano vigente
    "bimestre": "2",
    "unidade": dedup(u for u, _, _ in parsed),
    "objetos": dedup(sum((split_itens(o) for _, o, _ in parsed), [])),
    "habilidade": dedup(sum(((h.split(", ") if h else []) for _, _, h in parsed), [])),
    "estrategias": "Discussão em grupo / roda de conversa; Seminário; Trabalho em duplas ou trios; Outra estratégia.\nAtravés de vídeos do material digital, trabalhos apresentados pelos alunos e convite a professores de academias de lutas para uma aula prática.",
    "recursos": "Slides do material digital; Vídeo; Computador / notebook",
    "avaliacao": "Através de atividades elaboradas e executadas pelos alunos.",
    "adaptacao": "",
    "verificacao": "Observação do desempenho em atividades; Trabalho em grupo; Outro instrumento",
}

if __name__ == "__main__":
    saida = Path(__file__).resolve().parent.parent / "EXEMPLO_Plano_de_Ensino_v2.docx"
    gerar(str(saida), dados)
