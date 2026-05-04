import os

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from services.avaliador_frequencia import AvaliadorFrequencia
from services.configuracao import Configuracao
from services.periodo_letivo import garantir_bimestre_operacional
from services.runtime_paths import data_dir


class GeradorRelatorioProfessores:
    @staticmethod
    def _nome_aluno(aluno):
        return aluno.nome.title()

    @staticmethod
    def _ajustes_por_disciplina(aluno, bimestre):
        return getattr(aluno, "ajustes_medias_conselho", {}).get(bimestre, {})

    @staticmethod
    def _formatar_media(valor):
        return "-" if valor is None else f"{float(valor):.1f}"

    @staticmethod
    def _adicionar_tabela(doc, titulo, colunas, linhas):
        doc.add_paragraph()
        bloco = doc.add_paragraph()
        bloco.add_run(titulo).bold = True

        if not linhas:
            doc.add_paragraph("Nenhum aluno.")
            return

        tabela = doc.add_table(rows=1, cols=len(colunas))
        tabela.style = "Table Grid"
        tabela.autofit = True

        for idx, coluna in enumerate(colunas):
            cell = tabela.rows[0].cells[idx]
            cell.text = coluna
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragrafo in cell.paragraphs:
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragrafo.runs:
                    run.bold = True

        for linha in linhas:
            row = tabela.add_row().cells
            for idx, valor in enumerate(linha):
                row[idx].text = str(valor)
                row[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    @staticmethod
    def gerar(turma, bimestre, caminho_saida=None):
        bimestre = garantir_bimestre_operacional(bimestre)
        doc = Document()

        # Fonte padrão
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)

        titulo = doc.add_paragraph()
        run = titulo.add_run(
            f"Relatório Pedagógico – Bimestre {bimestre} – Turma {turma.codigo}"
        )
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()

        # ===================== AGRUPADO POR DISCIPLINA =====================
        nota_minima = Configuracao.obter_nota_minima()

        carga = turma.carga_horaria.get(bimestre, {})
        por_disciplina_defasagem = {}
        por_disciplina_ajustes = {}
        por_disciplina_faltas = {}
        encontrou_medias = False

        for aluno in turma.alunos.values():
            if not getattr(aluno, "ativo", True):
                continue

            medias_bim = getattr(aluno, "medias", {}).get(bimestre, {})
            ajustes_bim = GeradorRelatorioProfessores._ajustes_por_disciplina(aluno, bimestre)
            disciplinas_ajustes = set(ajustes_bim.keys())
            for disciplina in sorted(set(medias_bim.keys()) | disciplinas_ajustes):
                media = medias_bim.get(disciplina)
                if media is not None:
                    encontrou_medias = True
                nome_fmt = GeradorRelatorioProfessores._nome_aluno(aluno)
                ajuste = ajustes_bim.get(disciplina)
                if ajuste and ajuste.get("media_ajustada") is not None:
                    por_disciplina_ajustes.setdefault(disciplina, []).append(
                        (
                            nome_fmt,
                            ajuste.get("media_original", media),
                            ajuste.get("media_ajustada"),
                            ajuste.get("observacao", "").strip(),
                        )
                    )
                elif media is not None and media < nota_minima:
                    por_disciplina_defasagem.setdefault(disciplina, []).append((nome_fmt, media))

            faltas_bim = aluno.frequencia.get(bimestre, {})
            for disciplina, faltas in faltas_bim.items():
                total_aulas = carga.get(disciplina)
                if not total_aulas or total_aulas <= 0:
                    continue
                if (faltas / total_aulas) > AvaliadorFrequencia.LIMITE:
                    percentual = (faltas / total_aulas) * 100
                    nome_fmt = GeradorRelatorioProfessores._nome_aluno(aluno)
                    por_disciplina_faltas.setdefault(disciplina, []).append(
                        (nome_fmt, faltas, total_aulas, percentual)
                    )

        # Disciplinas a considerar (união de ajustes, notas e faltas)
        disciplinas = sorted(
            set(por_disciplina_defasagem.keys())
            | set(por_disciplina_ajustes.keys())
            | set(por_disciplina_faltas.keys())
        )

        if not encontrou_medias:
            doc.add_paragraph(
                "Nenhuma média encontrada para este bimestre. "
                "Reimporte o mapão para registrar as médias."
            )
            doc.add_paragraph()

        if not disciplinas:
            doc.add_paragraph("Nenhum registro encontrado para este bimestre.")
        else:
            for idx, disciplina in enumerate(disciplinas):
                doc.add_paragraph()
                cabecalho = doc.add_paragraph()
                cabecalho.add_run(f"Disciplina: {disciplina}").bold = True
                doc.add_paragraph("Tarefas do professor para registro e acompanhamento.")

                lista_ajustes = sorted(por_disciplina_ajustes.get(disciplina, []), key=lambda x: x[0])
                lista_faltas = sorted(
                    por_disciplina_faltas.get(disciplina, []), key=lambda x: x[0]
                )
                lista_notas = sorted(
                    por_disciplina_defasagem.get(disciplina, []), key=lambda x: x[0]
                )

                tarefas = []
                if lista_ajustes:
                    tarefas.append("Ajustar notas na Sala do Futuro para os alunos listados abaixo.")
                if lista_faltas:
                    tarefas.append("Organizar compensação de faltas para os alunos listados abaixo.")
                if lista_notas:
                    tarefas.append("Acompanhar a defasagem de nota dos alunos sem ajuste registrado.")
                if not tarefas:
                    tarefas.append("Nenhuma ação pendente para esta disciplina.")

                for numero, tarefa in enumerate(tarefas, start=1):
                    doc.add_paragraph(f"{numero}. {tarefa}")

                GeradorRelatorioProfessores._adicionar_tabela(
                    doc,
                    "Ajustar notas na Sala do Futuro",
                    ("Aluno", "Media original", "Media ajustada", "Observacao"),
                    [
                        (
                            nome,
                            GeradorRelatorioProfessores._formatar_media(media_original),
                            GeradorRelatorioProfessores._formatar_media(media_ajustada),
                            observacao or "-",
                        )
                        for nome, media_original, media_ajustada, observacao in lista_ajustes
                    ],
                )

                GeradorRelatorioProfessores._adicionar_tabela(
                    doc,
                    "Compensar faltas",
                    ("Aluno", "Faltas", "Aulas", "% Faltas"),
                    [
                        (nome, faltas, total, f"{percentual:.1f}%")
                        for nome, faltas, total, percentual in lista_faltas
                    ],
                )

                GeradorRelatorioProfessores._adicionar_tabela(
                    doc,
                    "Alunos com defasagem de nota sem ajuste",
                    ("Aluno", "Media atual"),
                    [(nome, f"{media:.1f}") for nome, media in lista_notas],
                )

                if idx < len(disciplinas) - 1:
                    doc.add_page_break()

        # ===================== SALVAR =====================
        if caminho_saida:
            pasta = os.path.dirname(caminho_saida)
            if pasta:
                os.makedirs(pasta, exist_ok=True)
            caminho = caminho_saida
        else:
            pasta = data_dir("relatorios")
            os.makedirs(pasta, exist_ok=True)
            caminho = os.path.join(
                pasta,
                f"relatorio_professores_{turma.codigo}_bim_{bimestre}.docx"
            )

        doc.save(caminho)
        return caminho
