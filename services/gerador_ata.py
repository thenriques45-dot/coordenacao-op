import os
from datetime import date, datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from services.configuracao import Configuracao
from services.preparador_ata import PreparadorAta
from services.periodo_letivo import garantir_bimestre_operacional


# ======================================================
# ABREVIAÇÕES DAS DISCIPLINAS
# ======================================================
ABREVIACOES = {
    "BIOLOGIA": "BIO",
    "FÍSICA": "FIS",
    "FISICA": "FIS",
    "GEOGRAFIA": "GEO",
    "HISTÓRIA": "HIST",
    "HISTORIA": "HIST",
    "LINGUA PORTUGUESA": "PORT",
    "MATEMATICA": "MAT",
    "QUIMICA": "QUI",
    "REDAÇÃO E LEITURA": "RED",
    "ARTE": "ART",
    "ARTE E MÍDIAS DIGITAIS": "ART",
    "EDUCACAO FISICA": "EDF",
    "FILOSOFIA E SOCIEDADE MODERNA": "FIL",
    "GEOPOLITICA": "GEOP",
    "LINGUA INGLESA": "ING",
    "PROJETO DE VIDA": "PV",
    "EDUCAÇÃO FINANCEIRA": "EFIN",
    "TECNOLOGIA E INOVAÇÃO": "TEC",
    "CIENCIAS": "CIE",
}


# ======================================================
# UTILIDADES VISUAIS
# ======================================================
def repetir_cabecalho(linha):
    tr = linha._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def pintar(celula, cor_hex):
    tc = celula._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), cor_hex)
    tcPr.append(shd)

def impedir_quebra_linha(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)

def remover_bordas(tabela):
    tbl = tabela._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')

    for edge in ('top','left','bottom','right','insideH','insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'nil')
        borders.append(elem)

    tblPr.append(borders)


# ======================================================
# DATA POR EXTENSO
# ======================================================
def data_por_extenso(data: date):
    meses = [
        "janeiro","fevereiro","março","abril","maio","junho",
        "julho","agosto","setembro","outubro","novembro","dezembro"
    ]

    nums = {
        1:"um",2:"dois",3:"três",4:"quatro",5:"cinco",6:"seis",7:"sete",8:"oito",9:"nove",
        10:"dez",11:"onze",12:"doze",13:"treze",14:"quatorze",15:"quinze",16:"dezesseis",
        17:"dezessete",18:"dezoito",19:"dezenove",20:"vinte",21:"vinte e um",22:"vinte e dois",
        23:"vinte e três",24:"vinte e quatro",25:"vinte e cinco",26:"vinte e seis",
        27:"vinte e sete",28:"vinte e oito",29:"vinte e nove",30:"trinta",31:"trinta e um"
    }

    return f"{nums[data.day]} de {meses[data.month-1]} de dois mil e {nums.get(data.year-2000,data.year)}"


# ======================================================
# GERADOR
# ======================================================
class GeradorAta:

    @staticmethod
    def gerar(
        turma,
        bimestre,
        data_conselho=None,
        confirmar_continuacao=None,
        log=None
    ):
        bimestre = garantir_bimestre_operacional(bimestre)
        if log is None:
            log = print

        doc = Document()

        # ======================================================
        # MARGENS
        # ======================================================
        for s in doc.sections:
            s.top_margin = Cm(1)
            s.bottom_margin = Cm(1)
            s.left_margin = Cm(1)
            s.right_margin = Cm(1)

        # ======================================================
        # FONTE GLOBAL
        # ======================================================
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10)

        if "TextoAta" not in doc.styles:
            estilo = doc.styles.add_style("TextoAta", WD_STYLE_TYPE.PARAGRAPH)
            estilo.font.name = "Calibri"
            estilo.font.size = Pt(10)

        # ======================================================
        # CABEÇALHO IMAGEM
        # ======================================================
        header = doc.sections[0].header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(
            "dados/imagens/cabecalho.jpg",
            width=Cm(12)
        )

        # ======================================================
        # TÍTULO
        # ======================================================
        doc.add_paragraph()
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r = titulo.add_run(f"CONSELHO DE CLASSE - {bimestre}º BIM/{turma.ano}")
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(128, 0, 128)

        # ======================================================
        # TEXTO INTRODUTÓRIO COMPLETO (RESTaurado)
        # ======================================================
        if data_conselho is None:
            entrada = input("Informe a data do Conselho (DD/MM/AAAA) ou Enter para hoje: ").strip()
            data_conselho = datetime.strptime(entrada, "%d/%m/%Y").date() if entrada else date.today()

        total = len(turma.alunos)
        frequentes = sum(1 for a in turma.alunos.values() if a.ativo)
        direcao_nome, direcao_pronome = Configuracao.obter_direcao()
        artigo = "da" if direcao_pronome == "F" else "do"
        titulo_direcao = "Diretora Sra." if direcao_pronome == "F" else "Diretor Sr."
        ciclo = getattr(turma,"ciclo","")

        intro = doc.add_paragraph(style="TextoAta")
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        intro.add_run(
            f"Aos {data_por_extenso(data_conselho)}, reuniram-se presencialmente a presidência {artigo} "
            f"{titulo_direcao} {direcao_nome}, equipe gestora, professores, estudantes e responsáveis da turma do "
            f"{turma.serie or ''} {turma.codigo} {ciclo} para procederem ao CONSELHO DE CLASSE. "
            f"Na abertura a diretora pautou que no conselho de classe devem ser colocadas situações que mereçam "
            f"um estudo de caso e registro de alternativas para intervenções pedagógicas que tenham como meta "
            f"o desenvolvimento do processo ensino/aprendizagem dos alunos. Foram tratados também os seguintes assuntos: "
            f"(1) levantamento de estudantes que não realizaram nenhuma das atividades e projetos; "
            f"(2) levantamento de estudantes que necessitam de compensação de ausência; "
            f"(3) estudantes com defasagem de habilidades e conteúdos para a respectiva série que necessitam de acompanhamento pedagógico; "
            f"(4) levantamento de estudantes que necessitam de recuperação e aprofundamento. "
            f"Para efeito de registro documental, verificou-se que a turma é composta por {total} estudantes matriculados, "
            f"sendo {frequentes} alunos frequentes, e destes estudantes frequentes não alcançaram a menção mínima nas disciplinas:"
        )

        # ======================================================
        # TABELA
        # ======================================================
        disciplinas = PreparadorAta.levantar_disciplinas(turma)
        alunos = PreparadorAta.preparar_alunos(turma, bimestre)

        # ======================================================
        # VALIDAÇÕES ANTES DE GERAR ATA
        # ======================================================

        # 1️⃣ mapão não importado (erro fatal)
        if bimestre not in turma.carga_horaria:
            raise ValueError(
                f"Mapão do {bimestre}º bimestre não foi importado.\n"
                "Importe o mapão antes de gerar a ata."
            )

        # 2️⃣ disciplinas vazias (erro fatal)
        if not disciplinas:
            raise ValueError(
                "Nenhuma disciplina encontrada.\n"
                "Verifique se o mapão foi importado corretamente."
            )

        # 3️⃣ alunos sem frequência (aviso + relatório)
        sem_freq = [
            a["nome"]
            for a in alunos
            if not a["linha_amarela"] and not a["frequencia_percentual"]
        ]

        if sem_freq:

            log("\n⚠ ATENÇÃO: alunos sem frequência importada:\n")
            for nome in sem_freq:
                log(" -", nome)

            # ---------- gera relatório ----------
            pasta_rel = "dados/relatorios"
            os.makedirs(pasta_rel, exist_ok=True)

            caminho_rel = os.path.join(
                pasta_rel,
                f"faltando_frequencia_{turma.codigo}_bim_{bimestre}.txt"
            )

            with open(caminho_rel, "w", encoding="utf-8") as f:
                f.write("ALUNOS SEM FREQUÊNCIA IMPORTADA\n\n")
                for nome in sem_freq:
                    f.write(nome + "\n")

            log(f"\nRelatório salvo em: {caminho_rel}\n")

            if confirmar_continuacao is not None:
                if not confirmar_continuacao(sem_freq, caminho_rel):
                    log("Geração da ata cancelada.")
                    return None
            else:
                # ---------- pergunta continuar ----------
                resp = input("Deseja continuar mesmo assim? (S/N): ").strip().upper()

                if resp != "S":
                    log("Geração da ata cancelada.")
                    return None

        cabecalhos = (
            ["Nº","ALUNO","STATUS"]
            + [ABREVIACOES.get(d.upper(), d[:4].upper()) for d in disciplinas]
            + ["FREQ (%)","ENCAM."]
        )

        tabela = doc.add_table(rows=1, cols=len(cabecalhos))
        tabela.style = "Table Grid"
        tabela.autofit = True

        # ========= LARGURAS CORRETAS =========
        tabela.columns[1].width = Cm(6.0)


        # ========= CABEÇALHO =========
        row = tabela.rows[0]
        repetir_cabecalho(row)

        for i,t in enumerate(cabecalhos):
            cell = row.cells[i]
            pintar(cell,"E6E6E6")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(t.upper())
            r.bold = True
            r.font.size = Pt(8)

        # ========= LINHAS =========
        for aluno in alunos:

            linha_row = tabela.add_row()
            impedir_quebra_linha(linha_row)
            linha = linha_row.cells

            for c in linha:
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Nº
            p = linha[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.clear()
            p.add_run(str(aluno["numero"]))

            # Nome
            p = linha[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.clear()
            p.add_run(aluno["nome"])

            # Status
            p = linha[2].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.clear()
            p.add_run(aluno["status"] or "")

            idx = 3
            for d in disciplinas:
                p = linha[idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.clear()
                if d in aluno["defasagens"]:
                    p.add_run("X")
                idx += 1

            # Frequência
            p = linha[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.clear()
            freq_texto = aluno.get("frequencia_percentual", "")
            p.add_run(freq_texto if freq_texto not in (None, "") else "")
            idx += 1

            linha[idx].text = aluno.get("encaminhamento","")

            # fonte
            for c in linha:
                for r in c.paragraphs[0].runs:
                    r.font.size = Pt(7.5)

            pintar(linha[0],"E6E6E6")
            pintar(linha[1],"E6E6E6")

            if aluno["status"]:
                for c in linha:
                    pintar(c,"FFFF00")

        
        # ======================================================
        # LISTA FIXA DE ENCAMINHAMENTOS (ESTILIZADA)
        # ======================================================

        doc.add_paragraph()

        titulo_enc = doc.add_paragraph()
        titulo_enc.add_run("Outras observações e encaminhamentos:").bold = True

        textos = [
            "Dificuldade em ler, interpretar e associar dados, tabelas, figuras, produzir textos e resolver situações problemas",
            "Confrontar ideias e opiniões, manifestando-se de forma argumentativa",
            "Dedicar-se mais ao estudo em casa.",
            "Prestar mais atenção às explicações do professor, tirar dúvidas, realizar as tarefas em aula nos prazos estipulados",
            "Frequência às aulas.",
            "Acompanhar diariamente, dialogar e orientar o estudante sobre as atividades escolares",
            "Estabelecer horas de estudo em casa, incentivando o hábito de estudar",
            "Comparecer às reuniões e conversar com professores e coordenadores pedagógicos",
            "Recuperação contínua",
            "Tarefas auxiliares para superação das dificuldades específicas do estudante"
        ]

        tabela_enc = doc.add_table(rows=5, cols=4)
        tabela_enc.style = "Table Grid"
        tabela_enc.autofit = True
        tabela_enc.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 👉 números mais estreitos
        for row in tabela_enc.rows:
            row.cells[0].width = Cm(1.0)   # nº esquerdo
            row.cells[1].width = Cm(9.4)   # texto esquerdo
            row.cells[2].width = Cm(1.0)   # nº direito
            row.cells[3].width = Cm(9.4)   # texto direito
        
        cinza_nums = {1, 3, 5, 6, 8, 10}

        idx = 0
        for i in range(5):
            linha = tabela_enc.rows[i].cells

        # ---------- COLUNA ESQUERDA ----------
            numero = i + 1
            texto = textos[i]

            p = linha[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{numero}.")
            run.font.size = Pt(9)

            if numero in cinza_nums:
                pintar(linha[0], "E6E6E6")

            linha[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = linha[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(texto).font.size = Pt(9)
            linha[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


            # ---------- COLUNA DIREITA ----------
            numero = i + 6
            texto = textos[i + 5]

            p = linha[2].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{numero}.")
            run.font.size = Pt(9)

            if numero in cinza_nums:
                pintar(linha[2], "E6E6E6")

            linha[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = linha[3].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(texto).font.size = Pt(9)
            linha[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # ======================================================
        # ASSINATURA DOS PROFESSORES
        # ======================================================

        doc.add_paragraph()
        titulo_ass = doc.add_paragraph("ASSINATURA DOS PROFESSORES:")
        titulo_ass.runs[0].bold = True

        disciplinas_ass = sorted(disciplinas)

        colunas = 4
        linhas = (len(disciplinas_ass) + colunas - 1) // colunas

        tabela_ass = doc.add_table(rows=linhas, cols=colunas)
        tabela_ass.autofit = True

        remover_bordas(tabela_ass)

        idx = 0

        for r in range(linhas):
            for c in range(colunas):

                cell = tabela_ass.rows[r].cells[c]  # ✅ nome correto

                if idx < len(disciplinas_ass):      # ✅ nome correto
                    p = cell.paragraphs[0]

                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.line_spacing = 1.5

                    run = p.add_run(disciplinas_ass[idx])  # ✅ nome correto
                    run.font.size = Pt(9)

                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    idx += 1

        # ======================================================
        # ASSINATURAS FINAIS (SEM TABELA)
        # ======================================================

        doc.add_paragraph("\n")

        container = doc.add_table(rows=1, cols=2)
        remover_bordas(container)
        container.autofit = False
        container.columns[0].width = Cm(7)
        container.columns[1].width = Cm(7)

        titulos = ["Coordenação Pedagógica", "Direção"]

        for i, titulo in enumerate(titulos):

                cell = container.rows[0].cells[i]

                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)

                # linha de assinatura (underline fake)
                linha = p.add_run("______________________________")
                linha.font.size = Pt(12)

                p.add_run("\n")

                texto = p.add_run(titulo)
                texto.bold = True
                texto.font.size = Pt(12)


        # ======================================================
        # SALVAR
        # ======================================================
        pasta = "dados/atas"
        os.makedirs(pasta,exist_ok=True)

        caminho = os.path.join(
            pasta,
            f"ata_{turma.codigo}_bimestre_{bimestre}.docx"
        )

        doc.save(caminho)
        return caminho
