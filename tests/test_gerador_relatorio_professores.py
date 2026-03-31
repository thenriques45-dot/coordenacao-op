import os
import tempfile
import unittest
from unittest.mock import patch

from docx import Document

from domain.aluno import Aluno
from domain.turma import Turma
from services.gerador_relatorio_professores import GeradorRelatorioProfessores


class TestGeradorRelatorioProfessores(unittest.TestCase):
    def setUp(self):
        self._cwd = os.getcwd()
        self._tmpdir = tempfile.TemporaryDirectory()
        os.chdir(self._tmpdir.name)

    def tearDown(self):
        os.chdir(self._cwd)
        self._tmpdir.cleanup()

    def _texto_docx(self, caminho):
        doc = Document(caminho)
        textos = [p.text for p in doc.paragraphs]
        for tabela in doc.tables:
            for row in tabela.rows:
                for cell in row.cells:
                    textos.extend(p.text for p in cell.paragraphs)
        return "\n".join(t for t in textos if t)

    def test_ignora_alunos_inativos(self):
        turma = Turma("2A", 2026)
        turma.carga_horaria["1"] = {"MATEMATICA": 20}

        ativo = Aluno("1", "ALUNO ATIVO", ativo=True)
        ativo.medias["1"] = {"MATEMATICA": 4.0}
        ativo.frequencia["1"] = {"MATEMATICA": 6}

        inativo = Aluno("2", "ALUNO INATIVO", ativo=False)
        inativo.medias["1"] = {"MATEMATICA": 3.0}
        inativo.frequencia["1"] = {"MATEMATICA": 8}

        turma.adicionar_aluno(ativo)
        turma.adicionar_aluno(inativo)

        with patch("services.gerador_relatorio_professores.Configuracao.obter_nota_minima", return_value=5.0):
            caminho = GeradorRelatorioProfessores.gerar(turma, "1")

        texto = self._texto_docx(caminho)
        self.assertIn("Aluno Ativo", texto)
        self.assertNotIn("Aluno Inativo", texto)

    def test_mostra_aviso_quando_nao_ha_medias_validas(self):
        turma = Turma("2A", 2026)
        turma.carga_horaria["1"] = {"MATEMATICA": 20}

        aluno = Aluno("1", "ALUNO SEM MEDIA", ativo=True)
        aluno.medias["1"] = {"MATEMATICA": None}
        aluno.frequencia["1"] = {"MATEMATICA": 0}
        turma.adicionar_aluno(aluno)

        with patch("services.gerador_relatorio_professores.Configuracao.obter_nota_minima", return_value=5.0):
            caminho = GeradorRelatorioProfessores.gerar(turma, "1")

        texto = self._texto_docx(caminho)
        self.assertIn("Nenhuma média encontrada para este bimestre.", texto)

    def test_lista_apenas_disciplinas_com_excesso_de_faltas(self):
        turma = Turma("2A", 2026)
        turma.carga_horaria["1"] = {"MATEMATICA": 20, "HISTORIA": 20}

        aluno = Aluno("1", "ALUNO TESTE", ativo=True)
        aluno.medias["1"] = {"MATEMATICA": 7.0, "HISTORIA": 7.0}
        aluno.frequencia["1"] = {"MATEMATICA": 6, "HISTORIA": 4}
        turma.adicionar_aluno(aluno)

        with patch("services.gerador_relatorio_professores.Configuracao.obter_nota_minima", return_value=5.0):
            caminho = GeradorRelatorioProfessores.gerar(turma, "1")

        texto = self._texto_docx(caminho)
        self.assertIn("Compensar faltas", texto)
        self.assertIn("Aluno Teste", texto)
        self.assertIn("30.0%", texto)
        self.assertNotIn("20.0%", texto)

    def test_relatorio_destaca_ajuste_de_nota_com_observacao(self):
        turma = Turma("2A", 2026)
        turma.carga_horaria["1"] = {"MATEMATICA": 20}

        aluno = Aluno("1", "ALUNO TESTE", ativo=True)
        aluno.medias["1"] = {"MATEMATICA": 4.0}
        aluno.frequencia["1"] = {"MATEMATICA": 0}
        aluno.ajustes_medias_conselho = {
            "1": {
                "MATEMATICA": {
                    "media_original": 4.0,
                    "media_ajustada": 5.5,
                    "observacao": "Ajustar apos recuperacao paralela",
                }
            }
        }
        turma.adicionar_aluno(aluno)

        with patch("services.gerador_relatorio_professores.Configuracao.obter_nota_minima", return_value=5.0):
            caminho = GeradorRelatorioProfessores.gerar(turma, "1")

        texto = self._texto_docx(caminho)
        self.assertIn("Ajustar notas na Sala do Futuro", texto)
        self.assertIn("Aluno Teste", texto)
        self.assertIn("5.5", texto)
        self.assertIn("Ajustar apos recuperacao paralela", texto)

    def test_relatorio_separa_defasagem_sem_ajuste(self):
        turma = Turma("2A", 2026)
        turma.carga_horaria["1"] = {"MATEMATICA": 20}

        aluno = Aluno("1", "ALUNO DEFASAGEM", ativo=True)
        aluno.medias["1"] = {"MATEMATICA": 4.0}
        aluno.frequencia["1"] = {"MATEMATICA": 0}
        turma.adicionar_aluno(aluno)

        with patch("services.gerador_relatorio_professores.Configuracao.obter_nota_minima", return_value=5.0):
            caminho = GeradorRelatorioProfessores.gerar(turma, "1")

        texto = self._texto_docx(caminho)
        self.assertIn("Alunos com defasagem de nota sem ajuste", texto)
        self.assertIn("Aluno Defasagem", texto)
        self.assertIn("4.0", texto)


if __name__ == "__main__":
    unittest.main()
